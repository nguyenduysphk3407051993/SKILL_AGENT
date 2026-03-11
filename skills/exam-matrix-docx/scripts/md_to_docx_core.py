#!/usr/bin/env python3
import argparse
import os
import re
import traceback
import pypandoc
from docx import Document
from docx.shared import RGBColor, Pt, Cm
from docx.oxml.ns import qn
from docx.oxml.shared import OxmlElement
from docx.enum.text import WD_TAB_ALIGNMENT

BLUE_HEX = RGBColor(0, 0, 255)
TARGET_FONT = 'Times New Roman'
TARGET_SIZE = Pt(12)
TAB_PLACEHOLDER = '::::'


def normalize_markdown_table(table_str: str) -> str:
    lines = table_str.strip().split('\n')
    if not lines:
        return table_str
    row_cells_counts = []
    for line in lines:
        pipes = len(re.findall(r'(?<!\\)\|', line))
        row_cells_counts.append(pipes - 1 if pipes > 1 else 0)
    if not row_cells_counts:
        return table_str
    max_cols = max(row_cells_counts)
    fixed_lines = []
    for line in lines:
        pipes = len(re.findall(r'(?<!\\)\|', line))
        current_cols = pipes - 1
        if current_cols < max_cols and pipes > 0:
            missing = max_cols - current_cols
            if re.search(r'\|[\s:-]+\|', line):
                line = line.strip() + ' --- |' * missing
            else:
                line = line.strip() + ' |' * missing
        fixed_lines.append(line)
    return '\n'.join(fixed_lines)


def mask_tables(text: str):
    table_store = {}
    lines = text.split('\n')
    new_lines, buffer_table = [], []
    is_table_mode = False
    table_counter = 0

    for line in lines:
        stripped = line.strip()
        is_row = ('|' in line) and (len(stripped) > 2) and (not stripped.startswith('>'))
        if is_row:
            is_table_mode = True
            buffer_table.append(line)
        else:
            if is_table_mode:
                if buffer_table:
                    full_table_str = '\n'.join(buffer_table)
                    if re.search(r'\|[:\s-]{3,}\|', full_table_str):
                        normalized_table = normalize_markdown_table(full_table_str)
                        placeholder = f'{{{{__TABLE_BLOCK_{table_counter}__}}}}'
                        table_store[placeholder] = normalized_table
                        new_lines += ['', placeholder, '']
                        table_counter += 1
                    else:
                        new_lines.extend(buffer_table)
                buffer_table = []
                is_table_mode = False
            new_lines.append(line)

    if buffer_table:
        full_table_str = '\n'.join(buffer_table)
        if re.search(r'\|[:\s-]{3,}\|', full_table_str):
            normalized_table = normalize_markdown_table(full_table_str)
            placeholder = f'{{{{__TABLE_BLOCK_{table_counter}__}}}}'
            table_store[placeholder] = normalized_table
            new_lines += ['', placeholder, '']
        else:
            new_lines.extend(buffer_table)

    return '\n'.join(new_lines), table_store


def unmask_tables(text: str, table_store: dict):
    for placeholder, original_content in table_store.items():
        text = text.replace(placeholder, original_content)
    return text


def reformat_quiz_options(text: str) -> str:
    lines = text.split('\n')
    new_lines = []
    i = 0
    option_pattern = re.compile(r'^\s*([A-Da-d])([\.\)])\s*(.*)')
    THRESHOLD_1_4 = 22
    THRESHOLD_1_2 = 45

    while i < len(lines):
        line = lines[i]
        m_a = option_pattern.match(line)
        if m_a and m_a.group(1).lower() == 'a':
            first_char = m_a.group(1)
            is_lower = first_char.islower()
            expected = ['b', 'c', 'd'] if is_lower else ['B', 'C', 'D']
            current_batch = [m_a]
            temp_i = i + 1
            found_all = True
            for char in expected:
                while temp_i < len(lines) and not lines[temp_i].strip():
                    temp_i += 1
                if temp_i < len(lines):
                    m = option_pattern.match(lines[temp_i])
                    if m and m.group(1) == char:
                        current_batch.append(m)
                        temp_i += 1
                    else:
                        found_all = False
                        break
                else:
                    found_all = False
                    break

            if found_all:
                formatted_opts, raw_contents = [], []
                for m in current_batch:
                    code, sep, content = m.group(1), m.group(2), m.group(3).strip()
                    raw_contents.append(content)
                    formatted_opts.append(f'**{code}{sep}** {content}')
                max_len = max(len(c) for c in raw_contents) if raw_contents else 0
                if max_len < THRESHOLD_1_4:
                    new_lines.append(f'{formatted_opts[0]}{TAB_PLACEHOLDER}{formatted_opts[1]}{TAB_PLACEHOLDER}{formatted_opts[2]}{TAB_PLACEHOLDER}{formatted_opts[3]}')
                elif max_len < THRESHOLD_1_2:
                    new_lines.append(f'{formatted_opts[0]}{TAB_PLACEHOLDER}{formatted_opts[1]}')
                    new_lines.append(f'{formatted_opts[2]}{TAB_PLACEHOLDER}{formatted_opts[3]}')
                else:
                    new_lines.extend(formatted_opts)
                new_lines.append('')
                i = temp_i
                continue

        m_any = option_pattern.match(line)
        if m_any:
            code, sep, content = m_any.group(1), m_any.group(2), m_any.group(3).strip()
            new_lines.append(f'**{code}{sep}** {content}')
        else:
            new_lines.append(line)
        i += 1

    return '\n'.join(new_lines)


def preprocess_markdown(md_text: str) -> str:
    safe, table_store = mask_tables(md_text)
    content = re.sub(r'\s+\*\*$', '', safe, flags=re.M)
    content = re.sub(r'([^\n])\n', r'\1 \n', content)
    keywords_split = r'((?:Bài|Câu|Ví dụ)\s*\d+[a-z]?\.?|(?:Lời giải|Đáp án)\s*:?)'
    content = re.sub(re.compile(keywords_split, re.IGNORECASE), r'\n\n\1', content)
    content = re.sub(r'(^|\s)([A-D]\s?[\)\.])', r'\n\2', content, flags=re.IGNORECASE | re.M)
    content = re.sub(re.compile(r'^((?:Bài|Câu|Ví dụ)\s*\d+[a-z]?\.?)', re.IGNORECASE | re.M), r'**\1**', content)
    content = re.sub(re.compile(r'^((?:Lời giải|Đáp án)\s*:?)', re.IGNORECASE | re.M), r'**\1**', content)
    content = reformat_quiz_options(content)
    content = unmask_tables(content, table_store)
    content = re.sub(r'\n{3,}', '\n\n', content)
    return content


def ensure_default_template(path: str):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = TARGET_FONT
    style.font.size = TARGET_SIZE
    style.element.rPr.rFonts.set(qn('w:eastAsia'), TARGET_FONT)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15
    doc.save(path)


def format_run(run):
    if '**' in run.text:
        run.text = run.text.replace('**', '')
    if run.text.strip().startswith('* '):
        run.text = run.text.replace('* ', '', 1)
    run.font.name = TARGET_FONT
    run.font.size = TARGET_SIZE
    r = run._element
    if r.rPr is None:
        r.get_or_add_rPr()
    r.rPr.rFonts.set(qn('w:eastAsia'), TARGET_FONT)
    if run.bold:
        run.font.color.rgb = BLUE_HEX


def postprocess_docx(tmp_docx: str, output_docx: str):
    doc = Document(tmp_docx)

    # Math font
    try:
        namespaces = {
            'm': 'http://schemas.openxmlformats.org/officeDocument/2006/math',
            'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
        }
        for math_r in doc.element.xpath('.//m:r', namespaces=namespaces):
            rPr = math_r.find(qn('w:rPr'))
            if rPr is None:
                rPr = OxmlElement('w:rPr')
                math_r.insert(0, rPr)
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is None:
                rFonts = OxmlElement('w:rFonts')
                rPr.append(rFonts)
            rFonts.set(qn('w:ascii'), TARGET_FONT)
            rFonts.set(qn('w:hAnsi'), TARGET_FONT)
            rFonts.set(qn('w:cs'), TARGET_FONT)
            rFonts.set(qn('w:eastAsia'), TARGET_FONT)
    except Exception:
        pass

    is_option_para_regex = re.compile(r'^(\*\*)?[a-dA-D][\.\)]')
    for para in doc.paragraphs:
        if TAB_PLACEHOLDER in para.text:
            cnt = para.text.count(TAB_PLACEHOLDER)
            for run in para.runs:
                if TAB_PLACEHOLDER in run.text:
                    run.text = run.text.replace(TAB_PLACEHOLDER, '\t')
            para.paragraph_format.tab_stops.clear_all()
            if cnt >= 3:
                para.paragraph_format.tab_stops.add_tab_stop(int(Cm(4.5).emu), WD_TAB_ALIGNMENT.LEFT)
                para.paragraph_format.tab_stops.add_tab_stop(int(Cm(9.0).emu), WD_TAB_ALIGNMENT.LEFT)
                para.paragraph_format.tab_stops.add_tab_stop(int(Cm(13.5).emu), WD_TAB_ALIGNMENT.LEFT)
            elif cnt == 2:
                para.paragraph_format.tab_stops.add_tab_stop(int(Cm(4.5).emu), WD_TAB_ALIGNMENT.LEFT)
                para.paragraph_format.tab_stops.add_tab_stop(int(Cm(9.0).emu), WD_TAB_ALIGNMENT.LEFT)
            elif cnt == 1:
                para.paragraph_format.tab_stops.add_tab_stop(int(Cm(9.0).emu), WD_TAB_ALIGNMENT.LEFT)
        elif is_option_para_regex.match(para.text.strip()):
            para.paragraph_format.tab_stops.clear_all()

        for run in para.runs:
            format_run(run)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        format_run(run)

    doc.save(output_docx)


def convert(input_path: str, output_path: str, template_path: str = ''):
    with open(input_path, 'r', encoding='utf-8') as f:
        md_text = f.read()

    md_processed = preprocess_markdown(md_text)

    if not template_path or not os.path.exists(template_path):
        template_path = output_path + '.default_template.docx'
        ensure_default_template(template_path)

    tmp_docx = output_path + '.tmp.docx'
    extra_args = [f'--reference-doc={template_path}'] if template_path else []

    pypandoc.convert_text(
        md_processed,
        'docx',
        format='markdown+grid_tables+pipe_tables+backtick_code_blocks',
        outputfile=tmp_docx,
        extra_args=extra_args,
    )

    postprocess_docx(tmp_docx, output_path)

    if os.path.exists(tmp_docx):
        os.remove(tmp_docx)
    if template_path.endswith('.default_template.docx') and os.path.exists(template_path):
        os.remove(template_path)


def main():
    ap = argparse.ArgumentParser(description='Convert markdown to docx with PyQt6-core-equivalent logic (CLI)')
    ap.add_argument('--input', required=True, help='Input markdown/txt file')
    ap.add_argument('--output', required=True, help='Output docx file')
    ap.add_argument('--template', default='', help='Optional reference docx template')
    args = ap.parse_args()

    try:
        convert(args.input, args.output, args.template)
        print(f'OK: {args.output}')
    except Exception as e:
        print('ERROR:', e)
        print(traceback.format_exc())
        raise


if __name__ == '__main__':
    main()
