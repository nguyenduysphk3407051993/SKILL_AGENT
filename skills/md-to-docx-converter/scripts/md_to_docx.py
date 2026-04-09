import sys
import os
import re
import argparse
import pypandoc
import docx
from docx import Document
from docx.shared import RGBColor, Pt, Cm
from docx.oxml.ns import qn
from docx.oxml.shared import OxmlElement
from docx.enum.text import WD_TAB_ALIGNMENT

class MarkdownToDocxConverter:
    def __init__(self, template_filepath=None):
        self.template_filepath = template_filepath
        self.table_store = {}
        
        # --- CẤU HÌNH ---
        self.BLUE_HEX = RGBColor(0, 0, 255) 
        self.TARGET_FONT = 'Times New Roman'
        self.TARGET_SIZE = Pt(12) 

    def normalize_markdown_table(self, table_str):
        """Chuẩn hóa bảng Markdown: đảm bảo số cột ở header, separator và body bằng nhau."""
        lines = table_str.strip().split('\n')
        if not lines:
            return table_str

        row_cells_counts = []
        for line in lines:
            pipes = len(re.findall(r'(?<!\\)\|', line))
            if pipes > 1:
                row_cells_counts.append(pipes - 1)
            else:
                row_cells_counts.append(0)

        if not row_cells_counts:
            return table_str

        max_cols = max(row_cells_counts)
        fixed_lines = []
        for line in lines:
            pipes = len(re.findall(r'(?<!\\)\|', line))
            current_cols = pipes - 1
            
            if current_cols < max_cols and pipes > 0:
                missing = max_cols - current_cols
                if re.search(r'\|[\s:-]+\|', line):
                    line = line.strip() + ' --- |' * missing
                else:
                    line = line.strip() + ' |' * missing
            fixed_lines.append(line)
            
        return "\n".join(fixed_lines)

    def mask_tables(self, text):
        """Giấu bảng Markdown để không bị Regex phá vỡ, đồng thời chuẩn hóa bảng"""
        self.table_store = {}
        lines = text.split('\n')
        new_lines = []
        buffer_table = []
        is_table_mode = False
        table_counter = 0

        for line in lines:
            stripped = line.strip()
            is_row = ('|' in line) and (len(stripped) > 2) and (not line.strip().startswith('>'))

            if is_row:
                is_table_mode = True
                buffer_table.append(line)
            else:
                if is_table_mode:
                    if buffer_table:
                        full_table_str = "\n".join(buffer_table)
                        if re.search(r'\|[:\s-]{3,}\|', full_table_str):
                            normalized_table = self.normalize_markdown_table(full_table_str)
                            placeholder = f"{{{{__TABLE_BLOCK_{table_counter}__}}}}"
                            self.table_store[placeholder] = normalized_table
                            new_lines.extend(["", placeholder, ""])
                            table_counter += 1
                        else:
                            new_lines.extend(buffer_table)
                        buffer_table = []
                    is_table_mode = False
                new_lines.append(line)

        if buffer_table:
            full_table_str = "\n".join(buffer_table)
            if re.search(r'\|[:\s-]{3,}\|', full_table_str):
                normalized_table = self.normalize_markdown_table(full_table_str)
                placeholder = f"{{{{__TABLE_BLOCK_{table_counter}__}}}}"
                self.table_store[placeholder] = normalized_table
                new_lines.extend(["", placeholder, ""])
            else:
                new_lines.extend(buffer_table)

        print(f"🔎 Đã tìm thấy và bảo vệ {len(self.table_store)} bảng dữ liệu.")
        return "\n".join(new_lines)

    def unmask_tables(self, text):
        for placeholder, original_content in self.table_store.items():
            text = text.replace(placeholder, original_content)
        return text

    def reformat_quiz_options(self, text):
        """Dàn trang trắc nghiệm: 1 dòng, 2 dòng, hoặc 4 dòng tùy độ dài."""
        print("-> Đang phân tích và dàn trang trắc nghiệm (1 dòng / 2 dòng / 4 dòng)...")
        lines = text.split('\n')
        new_lines = []
        i = 0
        
        TAB_PLACEHOLDER = "::::"
        option_pattern = re.compile(r'^\s*([A-D])([\.\)])\s*(.*)', re.IGNORECASE)

        THRESHOLD_1_4 = 22  # 1/4 dòng
        THRESHOLD_1_2 = 45  # 1/2 dòng

        while i < len(lines):
            line = lines[i]
            match_a = option_pattern.match(line)
            
            if match_a and match_a.group(1).lower() == 'a':
                first_char = match_a.group(1) 
                is_lower = first_char.islower()
                expected_seq = ['b', 'c', 'd'] if is_lower else ['B', 'C', 'D']
                
                current_batch = [match_a]
                temp_i = i + 1
                found_all = True
                
                for char in expected_seq:
                    while temp_i < len(lines) and not lines[temp_i].strip():
                        temp_i += 1
                    
                    if temp_i < len(lines):
                        m = option_pattern.match(lines[temp_i])
                        if m and m.group(1) == char:
                            current_batch.append(m)
                            temp_i += 1
                        else:
                            found_all = False
                            break
                    else:
                        found_all = False
                        break
                
                if found_all:
                    formatted_opts = []
                    raw_contents = []
                    for m in current_batch:
                        char_code = m.group(1)
                        sep = m.group(2)
                        content = m.group(3).strip()
                        raw_contents.append(content)
                        formatted_opts.append(f"**{char_code}{sep}** {content}")
                    
                    max_len = max(len(c) for c in raw_contents)
                    
                    if max_len < THRESHOLD_1_4:
                        # 1 dòng (4 cột)
                        t = f"{TAB_PLACEHOLDER}{formatted_opts[0]}{TAB_PLACEHOLDER}{formatted_opts[1]}{TAB_PLACEHOLDER}{formatted_opts[2]}{TAB_PLACEHOLDER}{formatted_opts[3]}"
                        new_lines.append(t)
                    elif max_len < THRESHOLD_1_2:
                        # 2 dòng (2 cột)
                        new_lines.extend([
                            f"{TAB_PLACEHOLDER}{formatted_opts[0]}{TAB_PLACEHOLDER}{formatted_opts[1]}",
                            f"{TAB_PLACEHOLDER}{formatted_opts[2]}{TAB_PLACEHOLDER}{formatted_opts[3]}"
                        ])
                    else:
                        # 4 dòng
                        for opt in formatted_opts:
                            new_lines.extend([f"{TAB_PLACEHOLDER}{opt}", ""])
                    i = temp_i
                    continue

            # Fallback
            match_any = option_pattern.match(line)
            if match_any:
                char_code = match_any.group(1)
                sep = match_any.group(2)
                content = match_any.group(3).strip()
                new_lines.append(f"{TAB_PLACEHOLDER}**{char_code}{sep}** {content}")
            else:
                new_lines.append(line)
            i += 1
            
        return "\n".join(new_lines)
    
    def process_markdown(self, markdown_content):
        # 1. Masking tables
        content_safe = self.mask_tables(markdown_content)
        
        # 2. Basic cleanup
        content_processed = re.sub(r'\s+\*\*$', '', content_safe, flags=re.M)
        content_processed = re.sub(r'([^\n])\n', r'\1  \n', content_processed) # Bảo toàn ngắt dòng

        # 3. Handle bold titles (Bài, Câu, Ví dụ, Lời giải, Đáp án)
        keywords_split = r'((?:Bài|Câu|Ví dụ)\s*\d+[a-z]?\.?|(?:Lời giải|Đáp án)\s*:?)'
        content_processed = re.sub(re.compile(keywords_split, re.IGNORECASE), r'\n\n\1', content_processed)
        content_processed = re.sub(r'(^|\s)([A-D]\s?[\)\.])', r'\n\2', content_processed, flags=re.IGNORECASE | re.M)
        
        pattern_bold_title = re.compile(r'^((?:Bài|Câu|Ví dụ)\s*\d+[a-z]?\.?)', re.IGNORECASE | re.M)
        content_processed = re.sub(pattern_bold_title, r'**\1**', content_processed)

        pattern_bold_sol = re.compile(r'^((?:Lời giải|Đáp án)\s*:?)', re.IGNORECASE | re.M)
        content_processed = re.sub(pattern_bold_sol, r'**\1**', content_processed)

        # 4. Reformat options
        content_processed = self.reformat_quiz_options(content_processed)
        
        # 5. Unmask tables & final cleanup
        content_final_md = self.unmask_tables(content_processed)
        content_final_md = re.sub(r'\n{3,}', '\n\n', content_final_md)
        return content_final_md

    def format_run(self, run):
        # 1. Cleaner
        if '**' in run.text:
            run.text = run.text.replace('**', '')
        if run.text.strip().startswith('* '):
            run.text = run.text.replace('* ', '', 1)
        
        # 2. Font configuration
        run.font.name = self.TARGET_FONT
        run.font.size = self.TARGET_SIZE
        r = run._element
        if r.rPr is None:
            r.get_or_add_rPr()
        r.rPr.rFonts.set(qn('w:eastAsia'), self.TARGET_FONT)
        
        # 3. Coloring bold elements
        if run.bold:
            run.font.color.rgb = self.BLUE_HEX

    def post_process_docx(self, docx_path):
        doc = Document(docx_path)
        
        print("-> Thiết lập Margins (L 2cm, R 1cm, T 1cm, B 1cm)...")
        for section in doc.sections:
            section.top_margin = Cm(1)
            section.bottom_margin = Cm(1)
            section.left_margin = Cm(2)
            section.right_margin = Cm(1)

        # --- Math Font Treatment ---
        print("-> Định dạng font Times New Roman cho công thức Toán...")
        namespaces = {
            'm': 'http://schemas.openxmlformats.org/officeDocument/2006/math',
            'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
        }
        for math_r in doc.element.xpath('.//m:r'):
            rPr = math_r.find(qn('w:rPr'))
            if rPr is None:
                rPr = OxmlElement('w:rPr')
                math_r.insert(0, rPr)
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is None:
                rFonts = OxmlElement('w:rFonts')
                rPr.append(rFonts)
            rFonts.set(qn('w:ascii'), self.TARGET_FONT)
            rFonts.set(qn('w:hAnsi'), self.TARGET_FONT)
            rFonts.set(qn('w:cs'), self.TARGET_FONT)
            rFonts.set(qn('w:eastAsia'), self.TARGET_FONT)

        # --- Formatting runs & Tab stops ---
        print("-> Thiết lập Tab Stop và xử lý định dạng...")
        TAB_PLACEHOLDER = "::::"
        is_option_para_regex = re.compile(r'^(\*\*)?[a-dA-D][\.\)]')

        for para in doc.paragraphs:
            if TAB_PLACEHOLDER in para.text:
                placeholder_count = para.text.count(TAB_PLACEHOLDER)
                for run in para.runs:
                    if TAB_PLACEHOLDER in run.text:
                        run.text = run.text.replace(TAB_PLACEHOLDER, "\t")
                
                para.paragraph_format.tab_stops.clear_all()
                para.paragraph_format.left_indent = Cm(0)
                para.paragraph_format.first_line_indent = Cm(0)
                
                if placeholder_count >= 4: 
                    para.paragraph_format.tab_stops.add_tab_stop(int(Cm(0.5).emu), WD_TAB_ALIGNMENT.LEFT)
                    para.paragraph_format.tab_stops.add_tab_stop(int(Cm(4.5).emu), WD_TAB_ALIGNMENT.LEFT)
                    para.paragraph_format.tab_stops.add_tab_stop(int(Cm(9.0).emu), WD_TAB_ALIGNMENT.LEFT)
                    para.paragraph_format.tab_stops.add_tab_stop(int(Cm(13.5).emu), WD_TAB_ALIGNMENT.LEFT)
                elif placeholder_count == 3: 
                    para.paragraph_format.tab_stops.add_tab_stop(int(Cm(0.5).emu), WD_TAB_ALIGNMENT.LEFT)
                    para.paragraph_format.tab_stops.add_tab_stop(int(Cm(6.0).emu), WD_TAB_ALIGNMENT.LEFT)
                    para.paragraph_format.tab_stops.add_tab_stop(int(Cm(12.0).emu), WD_TAB_ALIGNMENT.LEFT)
                elif placeholder_count == 2: 
                    para.paragraph_format.tab_stops.add_tab_stop(int(Cm(0.5).emu), WD_TAB_ALIGNMENT.LEFT)
                    para.paragraph_format.tab_stops.add_tab_stop(int(Cm(9.0).emu), WD_TAB_ALIGNMENT.LEFT)
                elif placeholder_count == 1: 
                    para.paragraph_format.tab_stops.add_tab_stop(int(Cm(0.5).emu), WD_TAB_ALIGNMENT.LEFT)
            
            elif is_option_para_regex.match(para.text.strip()):
                para.paragraph_format.tab_stops.clear_all()
                para.paragraph_format.left_indent = Cm(0)
                para.paragraph_format.first_line_indent = Cm(0)

            for run in para.runs:
                self.format_run(run)

        # Format tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            self.format_run(run)

        doc.save(docx_path)

    def convert(self, input_markdown_path, output_docx_path):
        print("--- Đọc file Markdown ---")
        with open(input_markdown_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        print("--- Xử lý nội dung Markdown ---")
        processed_md = self.process_markdown(content)
        
        # Save temporary
        temp_md = "temp_conversion.md"
        with open(temp_md, 'w', encoding='utf-8') as f:
            f.write(processed_md)
            
        print("--- Chạy Pandoc ---")
        extra_args = []
        if self.template_filepath and os.path.exists(self.template_filepath):
            extra_args = [f'--reference-doc={self.template_filepath}']
            
        try:
            pypandoc.convert_file(
                temp_md,
                'docx',
                format='markdown+grid_tables+pipe_tables+backtick_code_blocks',
                outputfile=output_docx_path,
                extra_args=extra_args
            )
            print("--- Hậu xử lý Docx ---")
            self.post_process_docx(output_docx_path)
            print(f"✅ Đã tạo thành công file: {output_docx_path}")
        except Exception as e:
            print(f"❌ Lỗi: {e}")
        finally:
            if os.path.exists(temp_md):
                os.remove(temp_md)

if __name__ == '__main__':
    parser = argparse.ArgumentParser(description="Chuyển đổi Markdown sang DOCX chuyên nghiệp.")
    parser.add_argument('-i', '--input', required=True, help="File markdown đầu vào")
    parser.add_argument('-o', '--output', required=True, help="File docx đầu ra")
    parser.add_argument('-t', '--template', required=False, help="File template docx tuỳ chọn")
    
    args = parser.parse_args()
    
    converter = MarkdownToDocxConverter(template_filepath=args.template)
    converter.convert(args.input, args.output)
